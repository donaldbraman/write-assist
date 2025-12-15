"""
Local file loading for source documents.
"""

import logging
from pathlib import Path

from write_assist.sources.models import SourceDocument, SourceLoadError, SourceType

logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst", ".text"}


def load_local_file(path: str) -> SourceDocument:
    """
    Load a local file as a source document.

    Args:
        path: Path to the local file

    Returns:
        SourceDocument with extracted content

    Raises:
        SourceLoadError: If file cannot be loaded
    """
    file_path = Path(path).expanduser().resolve()

    if not file_path.exists():
        raise SourceLoadError(path, "File not found")

    if not file_path.is_file():
        raise SourceLoadError(path, "Path is not a file")

    ext = file_path.suffix.lower()

    # Text-based files
    if ext in SUPPORTED_TEXT_EXTENSIONS:
        return _load_text_file(file_path, path)

    # PDF files
    if ext == ".pdf":
        return _load_pdf_file(file_path, path)

    # Word documents
    if ext in {".docx", ".doc"}:
        return _load_docx_file(file_path, path)

    # Fallback: try to read as text
    try:
        return _load_text_file(file_path, path)
    except UnicodeDecodeError as e:
        raise SourceLoadError(path, f"Unsupported file format: {ext}") from e


def _load_text_file(file_path: Path, original_path: str) -> SourceDocument:
    """Load a plain text file."""
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Try with latin-1 as fallback
        content = file_path.read_text(encoding="latin-1")

    return SourceDocument(
        source_type=SourceType.LOCAL_FILE,
        path=original_path,
        title=file_path.stem,
        content=content.strip(),
        word_count=len(content.split()),
        metadata={
            "extension": file_path.suffix,
            "size_bytes": file_path.stat().st_size,
        },
    )


def _load_pdf_file(file_path: Path, original_path: str) -> SourceDocument:
    """Load a PDF file (requires pypdf)."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise SourceLoadError(
            original_path, "PDF support requires 'pypdf' package: pip install pypdf"
        ) from e

    try:
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        content = "\n\n".join(pages)

        return SourceDocument(
            source_type=SourceType.LOCAL_FILE,
            path=original_path,
            title=file_path.stem,
            content=content.strip(),
            word_count=len(content.split()),
            metadata={
                "extension": ".pdf",
                "page_count": len(reader.pages),
                "size_bytes": file_path.stat().st_size,
            },
        )
    except Exception as e:
        raise SourceLoadError(original_path, f"Failed to read PDF: {e}") from e


def _load_docx_file(file_path: Path, original_path: str) -> SourceDocument:
    """Load a Word document (requires python-docx)."""
    try:
        from docx import Document
    except ImportError as e:
        raise SourceLoadError(
            original_path,
            "Word document support requires 'python-docx' package: pip install python-docx",
        ) from e

    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        return SourceDocument(
            source_type=SourceType.LOCAL_FILE,
            path=original_path,
            title=file_path.stem,
            content=content.strip(),
            word_count=len(content.split()),
            metadata={
                "extension": file_path.suffix,
                "paragraph_count": len(paragraphs),
                "size_bytes": file_path.stat().st_size,
            },
        )
    except Exception as e:
        raise SourceLoadError(original_path, f"Failed to read Word document: {e}") from e


def is_local_path(path: str) -> bool:
    """Check if a path looks like a local file path."""
    # Not a URL - could be absolute (/path/to/file) or relative (./file or file.txt)
    return not path.startswith(("http://", "https://", "ftp://"))
